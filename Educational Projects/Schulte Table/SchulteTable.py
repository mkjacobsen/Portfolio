# -*- coding: utf-8 -*-
"""
Created on Fri Jun 26 18:50:05 2020

@author: Matt Jacobsen
@title: Schulte Table Generator

Generates Schulte Table with randomized entries for worksheet printouts.
Schulte Tables aid in developing student concentration and focus.

Scroll down and enter the parameters for the Schulte Table in the 
ENTER PARAMETERS HERE block of code.
"""

# Required Packages
import os # Used to Navigate OS File Structure
from docx import Document
from docx.shared import Inches # Size control for defining Inches
from docx.shared import Pt # Point control for font size
from docx.enum.table import WD_TABLE_ALIGNMENT # Alignment for Table on Page
from docx.enum.text import WD_ALIGN_PARAGRAPH # Alignment for Text in Table
import random # Random generator or shuffler

def get_entries(size):
    """
    get_entries generates the list of entries for insertion into the table

    Parameters
    ----------
    size : int
        Defines the size of the square Schulte Table.

    Returns
    -------
    entries : list
        List of entries for insertion into Schulte table.

    """
    table_size = size*size  ## Determine the required number of entries
    i = 0 # Initialize the counter
    entries = [] # Initialize the entry list
    while i < table_size: # Loop until table size has been reached
        entries.append(i) # Enter the value of i into the table
        i+=1 # Increment the value of i
    
    return entries # Return the list of entries

def schulte_table(size, out_name, path):
    """
    schulte_table generates the Schulte table document and inserts the 
    values into the table structure. Then saves the document out.

    Parameters
    ----------
    size : int
        Size of square Schulte Table.
    out_name : str
        Output file name.
    path : str
        Output file path.

    Returns
    -------
    None.

    """
    entry_list = get_entries(size) # Generate the entries for the Schulte Table
    os.chdir(path) # Swap directories to path
    if out_name not in os.listdir(): # Check to see if the document already exists
        document = Document() # If not, start with an empty document
    else: 
        document = Document(out_name) # Otherwise, open the document
    
    # Add a Heading stating the size of the Schulte Table
    document.add_heading('Schulte Table: '+str(size)+' by '+str(size), level=1)
    random.shuffle(entry_list) # Shuffle the entry list
    table = document.add_table(rows=size, cols=size) # Add the table object to the document
    table.style = 'Table Grid' # Define the style of table
    table.autofit = False # Disable autofit
    
    # Initialize counters
    i_row = 0 
    i_col = 0
    index = 0
    
    for i_row in range(size): # Loop over all rows
        for i_col in range(size): # Loop over all columns
            table.cell(i_row,i_col).text = str(entry_list[index]) # Insert the value from the entry list
            table.cell(i_row,i_col).width = Inches(0.5) # Define the cell width
            index+=1 # Increment the entry index
            i_col+=1 # Increment the column index
        i_row+=1 # Increment the row index
    
    for row in table.rows: # Loop over rows
        for cell in row.cells: # Loop over cell in the row
            paragraphs = cell.paragraphs # Extract paragraphs (text entries)
            for paragraph in paragraphs: # For each one
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER # Aling the text center
                for run in paragraph.runs: # For each string of text
                    font = run.font # Access the font parameters
                    font.size = Pt(16) # Define the size of the font
    
    table.alignment = WD_TABLE_ALIGNMENT.CENTER # Align the table center
    document.save(out_name+'.docx')  # Save the document
#----------------------------------------------------------------------------#
########################### ENTER PARAMETERS HERE ############################
# Enter full file path for output document (leave the r in front of the '')
path = r''
# Enter the output file name
out_name = ''
# Enter the size of the Schulte Table you desire
size = 0
#----------------------------------------------------------------------------#
# Execute the script
if __name__=='__main__':
    schulte_table(size, out_name, path)
    print('Done!')
